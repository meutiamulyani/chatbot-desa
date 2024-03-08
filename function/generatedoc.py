from docx import Document
import provider.models

class Doc_Auto():
    # doc = Document()
    
    def __init__(self, db_con):
        # Create a new Document
        self.doc = Document()
        self.db_con = db_con
    
    # kondisi dokumen
    def doc_id(self, nomor_hp):
        # Heading judul form
        user = self.db_con.query(provider.models.user_activity).filter_by(no_hp=nomor_hp).first()
        if user:
            print('ngetes dulu la')
        else:
            print('kaga ada')
        return 
    
    def wrapper(self, user_id, jenis_form):
        if jenis_form == 'KTP':
            self.doc_id(user_id=user_id)
            print("lancar")
            ktp_doc = self.db_con.query(provider.models.form_ktp).filter_by(self.doc_id).first()
            self.doc.add_heading('SURAT KEPENGURUSAN KTP', level=1)
            self.doc.add_paragraph(f'Ujicoba dulu bosque {ktp_doc.id_user_activity}')
            self.doc.add_paragraph('Yaallah kenapa modulnya gak bisa')
            save = self.doc.save(f'dokumen{ktp_doc.id_user_activity}.docx')
            return save
        if jenis_form == 'Usaha':
            self.doc_usaha(user_id=user_id)
        if jenis_form == 'Rekomendasi':
            self.doc_rekomendasi(user_id=user_id)
        if jenis_form == 'SKTM':
            self.doc_sktm(user_id=user_id)
        return

    # Save the document
    # doc.save('my_document.docx')
