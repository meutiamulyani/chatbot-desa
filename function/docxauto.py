from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi import Request
import provider.models
from datetime import datetime, date

class Doc_Auto():
    # doc = Document()
    def __init__(self, db_con, model):
        # Create a new Document
        self.doc = Document()
        self.db_con = db_con
        self.model = model
    
    # kondisi dokumen
    def doc_id(self, nomor_hp):
        useract = self.db_con.query(self.model.user_activity).filter_by(no_hp=nomor_hp).first()
        print(useract)
        return useract

    def wrapper_doc(self, nomor_hp, jenis_form):
        # panggil doc_id untuk cek query nomor telpon dan activity
        user_activity = self.doc_id(nomor_hp=nomor_hp)
        print(f'user id: {nomor_hp}')

        # MEMBUAT KTP
        if jenis_form == 'KTP':
            print("lancar")
            # # data form_ktp
            ktp_doc = self.db_con.query(self.model.form_ktp).filter_by(id_user_activity = user_activity.id).first()
            title = self.doc.add_paragraph('FORMULIR PERMOHONAN KARTU TANDA PENDUDUK\nNOMOR SURAT: ............')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.doc.add_paragraph(
                'Yang bertanda tangan di bawah ini menerangkan dengan sebenar-benarnya bahwa: '
            )
            table = self.doc.add_table(rows=10, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Nama\t\t\t:'
            hdr_cells[1].text = 'Nomor KK\t\t:'
            hdr_cells[2].text = 'NIK\t\t\t:'
            hdr_cells[3].text = 'Alamat\t\t\t:'
            hdr_cells[4].text = 'RT/RW\t\t\t:'
            hdr_cells[5].text = 'Desa\t\t\t:'
            hdr_cells[6].text = 'Kecamatan\t\t:'
            hdr_cells[7].text = 'Kabupaten/Kota\t:'
            hdr_cells[8].text = 'Provinsi\t\t:'
            hdr_cells[9].text = 'Tujuan\t\t\t:'
                        
            for cell in table.columns[0].cells:
                cell.width = Inches(1.75)

            cols_cells = table.columns[1].cells
            cols_cells[0].text = ktp_doc.nama
            cols_cells[1].text = ktp_doc.no_kk
            cols_cells[2].text = ktp_doc.nik
            cols_cells[3].text = ktp_doc.alamat
            cols_cells[4].text = ktp_doc.rtrw
            cols_cells[5].text = ktp_doc.desa
            cols_cells[6].text = ktp_doc.kecamatan           
            cols_cells[7].text = ktp_doc.kabupaten_kota
            cols_cells[8].text = ktp_doc.provinsi
            cols_cells[9].text = ktp_doc.tujuan_surat

            self.doc.add_paragraph(
                f'Nama tersebut adalah benar warga {ktp_doc.desa}, dan yang bersangkutan mengajukan surat keterangan pengajuan KTP untuk keperluan {ktp_doc.tujuan_surat}'
            )
            self.doc.add_paragraph(
                f'Demikian surat keterangan ini dibuat, atas perhatian dan kerjasamanya terima kasih.'
            )

            today = date.today()
            uppercase_date = today.strftime("%B %d, %Y").upper()
            uppercase_desa = ktp_doc.kabupaten_kota.upper()
            ttd = self.doc.add_paragraph(f'{uppercase_desa}, {uppercase_date}\n\n\n(NAMA LENGKAP)')
            ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # rename doc
            current_time = datetime.now().strftime("%d-%m-%y_%H-%M-%S")
            file_name = f"ktp_{ktp_doc.nik}_{current_time}"

            # save doc
            save = self.doc.save(f'public/files/{file_name}.docx')
            return file_name

        elif jenis_form == 'Usaha':
            print("lancar")
            # # data form_usaha
            usaha_doc = self.db_con.query(self.model.form_usaha).filter_by(id_user_activity = user_activity.id).first()
            title = self.doc.add_paragraph('SURAT KETERANGAN USAHA\nNOMOR SURAT: ............')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.doc.add_paragraph(
                'Yang bertanda tangan di bawah ini Kepala Desa Dompyong Kulon menerangkan dengan sebenarnya bahwa: '
            )
            table = self.doc.add_table(rows=7, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Nama\t\t\t:'
            hdr_cells[1].text = 'NIK\t\t\t:'
            hdr_cells[2].text = 'Jenis Kelamin\t\t:'
            hdr_cells[3].text = 'Tempat/Tanggal Lahir\t:'
            hdr_cells[4].text = 'Alamat\t\t\t:'
            hdr_cells[5].text = 'Pekerjaan\t\t:'
            hdr_cells[6].text = 'RT/RW\t\t\t:'

            for cell in table.columns[0].cells:
                cell.width = Inches(1.75)

            cols_cells = table.columns[1].cells
            cols_cells[0].text = usaha_doc.nama
            cols_cells[1].text = usaha_doc.nik
            cols_cells[2].text = usaha_doc.jenis_kelamin
            cols_cells[3].text = usaha_doc.ttl
            cols_cells[4].text = usaha_doc.alamat
            cols_cells[5].text = usaha_doc.pekerjaan
            cols_cells[6].text = usaha_doc.rtrw           


            self.doc.add_paragraph(
                f'Sesuai dengan keterangan yang bersangkutan benar nama tersebut di atas mempunyai usaha sebagai berikut: '
            )

            # New table for usaha description
            table_usaha = self.doc.add_table(rows=4, cols=2)
            hdr_cells = table_usaha.columns[0].cells
            hdr_cells[0].text = 'Nama Usaha\t:'
            hdr_cells[1].text = 'Tanggal Usaha Dimulai:'
            hdr_cells[2].text = 'Alamat Usaha\t:'
            hdr_cells[3].text = 'Tujuan\t\t\t:'

            for cell in table_usaha.columns[0].cells:
                cell.width = Inches(1.75)

            cols_cells = table_usaha.columns[1].cells         
            cols_cells[0].text = usaha_doc.nama_usaha
            cols_cells[1].text = usaha_doc.start_usaha
            cols_cells[2].text = usaha_doc.alamat_usaha
            cols_cells[3].text = usaha_doc.tujuan_surat

            self.doc.add_paragraph(
                f'Demikian surat keterangan ini dibuat untuk dipergunakan sebagaimana mestinya.\n\n'
            )
            
            # rename doc
            current_time = datetime.now().strftime("%d-%m-%y_%H-%M-%S")
            file_name = f"usaha_{usaha_doc.nik}_{current_time}"
            today = date.today()

            uppercase_date = today.strftime("%B %d, %Y").upper()
            ttd = self.doc.add_paragraph(f'DOMPYONG KULON, {uppercase_date}\n\n\n(NAMA LENGKAP KEPALA DESA)')
            ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            save = self.doc.save(f'public/files/{file_name}.docx')
            return file_name

        elif jenis_form == 'Rekomendasi':
            print("lancar")
            # data form_usaha
            rec_doc = self.db_con.query(provider.models.form_rekomendasi).filter_by(id_user_activity = user_activity.id).first()
            
            paragraph = self.doc.add_paragraph('SURAT REKOMENDASI\nNOMOR SURAT: ............')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.doc.add_paragraph(
                'Yang bertanda tangan di bawah ini Kepala Desa Dompyong Kulon menerangkan dengan sebenarnya bahwa: '
            )

            table = self.doc.add_table(rows=7, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Nama\t\t\t:'
            hdr_cells[1].text = 'NIK\t\t\t:'
            hdr_cells[2].text = 'Tempat/Tanggal Lahir:'
            hdr_cells[3].text = 'Warganegara\t\t:'
            hdr_cells[4].text = 'Agama\t\t\t:'
            hdr_cells[5].text = 'Pekerjaan\t\t:'
            hdr_cells[6].text = 'Alamat\t\t\t:'
                        
            for cell in table.columns[0].cells:
                cell.width = Inches(1.75)

            cols_cells = table.columns[1].cells
            cols_cells[0].text = rec_doc.nama
            cols_cells[1].text = rec_doc.nik
            cols_cells[2].text = rec_doc.ttl
            cols_cells[3].text = rec_doc.warganegara
            cols_cells[4].text = rec_doc.agama
            cols_cells[5].text = rec_doc.pekerjaan
            cols_cells[6].text = rec_doc.alamat           
            
            self.doc.add_paragraph(
                f'Surat Rekomendasi ini dibuat untuk keperluan {rec_doc.tujuan_surat}.'
            )
            self.doc.add_paragraph(
                f'Demikian Rekomendasi ini dibuat dan diberikan kepada yang bersangkutan untuk dipergunakan sebagaimana mestinya.'
            )
            # rename doc
            current_time = datetime.now().strftime("%d-%m-%y_%H-%M-%S")
            file_name = f"rekomendasi_{rec_doc.nik}_{current_time}"
            today = date.today()

            uppercase_date = today.strftime("%B %d, %Y").upper()
            ttd = self.doc.add_paragraph(f'DOMPYONG KULON, {uppercase_date}\n\n\n(NAMA LENGKAP KEPALA DESA)')
            ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # generate docx
            save = self.doc.save(f'public/files/{file_name}.docx')
            return file_name

        elif jenis_form == 'SKTM':
            print("lancar")
            # data form_sktm
            sktm_doc = self.db_con.query(provider.models.form_sktm).filter_by(id_user_activity = user_activity.id).first()

            paragraph = self.doc.add_paragraph('SURAT KETERANGAN TIDAK MAMPU\nNOMOR SURAT: ............')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.doc.add_paragraph(
                'Yang bertanda tangan di bawah ini Kepala Desa Dompyong Kulon menerangkan dengan sebenarnya bahwa: '
            )

            table = self.doc.add_table(rows=9, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Nama\t\t\t:'
            hdr_cells[1].text = 'NIK\t\t\t:'
            hdr_cells[2].text = 'Jenis Kelamin\t\t:'
            hdr_cells[3].text = 'Tempat/Tanggal Lahir:'
            hdr_cells[4].text = 'Warganegara\t\t:'
            hdr_cells[5].text = 'Agama\t\t\t:'
            hdr_cells[6].text = 'Pekerjaan\t\t:'
            hdr_cells[7].text = 'Status\t\t\t:'
            hdr_cells[8].text = 'Alamat\t\t\t:'
                        
            for cell in table.columns[0].cells:
                cell.width = Inches(1.75)

            cols_cells = table.columns[1].cells
            cols_cells[0].text = sktm_doc.nama
            cols_cells[1].text = sktm_doc.nik
            cols_cells[2].text = sktm_doc.jenis_kelamin
            cols_cells[3].text = sktm_doc.ttl
            cols_cells[4].text = sktm_doc.warganegara
            cols_cells[5].text = sktm_doc.agama           
            cols_cells[6].text = sktm_doc.pekerjaan
            cols_cells[7].text = sktm_doc.status
            cols_cells[8].text = sktm_doc.alamat

            self.doc.add_paragraph(
                f'Surat ini dibuat untuk keperluan {sktm_doc.tujuan_surat}.\nDemikian SKTM ini dibuat dan diberikan kepada yang bersangkutan untuk dipergunakan sebagaimana mestinya.'
            )
            # rename doc
            current_time = datetime.now().strftime("%d-%m-%y_%H-%M-%S")
            file_name = f"sktm_{sktm_doc.nik}_{current_time}"
            today = date.today()

            uppercase_date = today.strftime("%B %d, %Y").upper()
            ttd = self.doc.add_paragraph(f'DOMPYONG KULON, {uppercase_date}\n\n\n(NAMA LENGKAP KEPALA DESA)')
            ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # generate docx
            save = self.doc.save(f'public/files/{file_name}.docx')
            return file_name

        else:
            print("error!")

        return None
