from docx import Document

# Create a new Document
doc = Document()

# Add a heading to the document
doc.add_heading('FORM KEHILANGAN KTP', level=1)

# Add paragraphs with text
doc.add_paragraph('This is a simple Word document created with python-docx.')
doc.add_paragraph('You can add more paragraphs as needed.')

# Save the document
doc.save('simple_document.docx')

print('Simple document created successfully.')
